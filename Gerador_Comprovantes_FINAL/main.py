import pandas as pd
from docx import Document
import re
import sys
import os
import subprocess
import tkinter as tk
from tkinter import messagebox
import unicodedata
from tkinter import ttk

# =============================
# CAMINHO
# =============================
def caminho_absoluto(nome_arquivo):
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, nome_arquivo)
    base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, nome_arquivo)

# =============================
# NORMALIZAR TEXTO
# =============================
def normalizar(texto):
    texto = str(texto).lower()
    texto = unicodedata.normalize("NFD", texto)
    texto = "".join(c for c in texto if unicodedata.category(c) != "Mn")
    return texto

# =============================
# IDENTIFICAR COLUNAS
# =============================
def identificar_colunas(colunas):
    nome_col = None
    cpf_col = None
    valor_col = None
    candidatos_valor = []

    for col in colunas:
        col_norm = normalizar(col)

        if not nome_col and ("nome" in col_norm or "bolsista" in col_norm):
            nome_col = col

        if not cpf_col and "cpf" in col_norm:
            cpf_col = col

        if "valor" in col_norm or "total" in col_norm or "recebido" in col_norm:
            score = 0
            if "valor total" in col_norm:
                score += 5
            if "total recebido" in col_norm:
                score += 5
            if "recebido" in col_norm:
                score += 3
            if "total" in col_norm:
                score += 2
            if "valor" in col_norm:
                score += 1
            if "parcela" in col_norm or "mensal" in col_norm:
                score -= 3

            candidatos_valor.append((score, col))

    if candidatos_valor:
        candidatos_valor.sort(reverse=True)
        valor_col = candidatos_valor[0][1]

    if not all([nome_col, cpf_col, valor_col]):
        raise Exception(
            "Não foi possível identificar automaticamente as colunas.\n\n"
            "Verifique se o Excel possui colunas de Nome, CPF e Valor."
        )

    return {"nome": nome_col, "cpf": cpf_col, "valor": valor_col}

# =============================
# FORMATADORES
# =============================
def formatar_cpf(cpf):
    cpf = re.sub(r"\D", "", str(cpf))
    return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}" if len(cpf) == 11 else cpf

def formatar_valor(valor):
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

# =============================
# SUBSTITUIR TEXTO NO WORD
# =============================
def substituir_texto(doc, mapa):
    for p in doc.paragraphs:
        for run in p.runs:
            for chave, valor in mapa.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, valor)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    for run in p.runs:
                        for chave, valor in mapa.items():
                            if chave in run.text:
                                run.text = run.text.replace(chave, valor)

# =============================
# CONVERTER DOCX → PDF
# =============================
LIBREOFFICE_PATHS = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/local/bin/soffice",
    "/usr/bin/soffice",
]

def encontrar_libreoffice():
    for path in LIBREOFFICE_PATHS:
        if os.path.exists(path):
            return path
    return None

def converter_para_pdf(docx_path, output_dir):
    soffice = encontrar_libreoffice()
    if not soffice:
        raise EnvironmentError(
            "LibreOffice não encontrado.\n\n"
            "Instale com:\n  brew install --cask libreoffice"
        )

    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
        capture_output=True,
        text=True,
        timeout=30
    )

    if result.returncode != 0:
        raise RuntimeError(f"Erro ao converter para PDF:\n{result.stderr}")

# =============================
# GERAR COMPROVANTES
# =============================
def gerar_comprovantes(barra, status_label, botao, janela):
    try:
        botao.config(state="disabled")
        status_label.config(text="Verificando LibreOffice...")

        # Valida LibreOffice antes de começar
        if not encontrar_libreoffice():
            raise EnvironmentError(
                "LibreOffice não encontrado.\n\n"
                "Instale com:\n  brew install --cask libreoffice"
            )

        xlsx_path = caminho_absoluto("dadosteste.xlsx")
        modelo_path = caminho_absoluto("MODELO COMPROVANTE RENDIMENTOS.docx")

        if not os.path.exists(xlsx_path):
            raise FileNotFoundError(f"Arquivo não encontrado:\n{xlsx_path}")
        if not os.path.exists(modelo_path):
            raise FileNotFoundError(f"Modelo não encontrado:\n{modelo_path}")

        df = pd.read_excel(xlsx_path)
        df.columns = df.columns.str.strip()

        colunas = identificar_colunas(df.columns)

        total = len(df)
        barra["maximum"] = total
        barra["value"] = 0

        base_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(base_dir, "output")
        temp_dir = os.path.join(base_dir, "output", "_temp_docx")
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(temp_dir, exist_ok=True)

        status_label.config(text="Iniciando...")

        for i, (_, linha) in enumerate(df.iterrows(), start=1):
            nome = str(linha[colunas["nome"]])
            cpf = formatar_cpf(linha[colunas["cpf"]])
            valor = formatar_valor(linha[colunas["valor"]])

            doc = Document(modelo_path)
            substituir_texto(doc, {
                "{{Nome}}": nome,
                "{{CPF}}": cpf,
                "{{Valor}}": valor
            })

            nome_arquivo = re.sub(r'[\\/*?:"<>|]', "_", nome)

            # Salva docx temporário
            docx_temp = os.path.join(temp_dir, f"{nome_arquivo}.docx")
            doc.save(docx_temp)

            # Converte para PDF
            status_label.config(text=f"Convertendo {i} de {total} para PDF...")
            janela.update_idletasks()
            converter_para_pdf(docx_temp, output_dir)

            # Remove docx temporário
            os.remove(docx_temp)

            barra["value"] = i
            status_label.config(text=f"Gerado {i} de {total}...")
            janela.update_idletasks()

        # Remove pasta temp
        os.rmdir(temp_dir)

        status_label.config(text="Concluído ✔")
        messagebox.showinfo(
            "Sucesso",
            f"PDFs gerados com sucesso!\n\nConfira a pasta:\n{output_dir}"
        )

    except Exception as e:
        messagebox.showerror("Erro", str(e))
        status_label.config(text="Erro ao gerar.")

    finally:
        botao.config(state="normal")

# =============================
# INTERFACE
# =============================
def main():
    janela = tk.Tk()
    janela.title("Gerador de Comprovantes")
    janela.geometry("420x280")
    janela.resizable(False, False)

    tk.Label(
        janela,
        text="Gerador de Comprovantes",
        font=("Segoe UI", 15, "bold")
    ).pack(pady=20)

    barra = ttk.Progressbar(
        janela,
        orient="horizontal",
        length=300,
        mode="determinate"
    )
    barra.pack(pady=15)

    status_label = tk.Label(
        janela,
        text="Aguardando...",
        font=("Segoe UI", 10)
    )
    status_label.pack()

    tk.Label(
        janela,
        text="Os comprovantes serão gerados em PDF",
        font=("Segoe UI", 9),
        fg="gray"
    ).pack(pady=(4, 0))

    botao = tk.Button(
        janela,
        text="Gerar comprovantes",
        font=("Segoe UI", 11),
        width=28,
        height=2,
        command=lambda: gerar_comprovantes(barra, status_label, botao, janela)
    )
    botao.pack(pady=16)

    janela.mainloop()

if __name__ == "__main__":
    main()